import streamlit as st
import google.generativeai as genai
import pdfplumber
import docx  # Used for reading and creating .docx files
import io
import os
from dotenv import load_dotenv

# --- Load environment variables ---
load_dotenv()

# --- Page Configuration ---
st.set_page_config(
    page_title="Professional Translation Workflow",
    page_icon="🌐",
    layout="wide",
)

# --- Helper Functions ---

def read_file(uploaded_file):
    """Reads the content of an uploaded file (txt, pdf, docx)."""
    try:
        file_type = uploaded_file.type
        file_name = uploaded_file.name

        if file_type == "text/plain":
            return uploaded_file.getvalue().decode("utf-8")
        
        elif file_type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        
        else:
            st.error(f"Unsupported file format: {file_type} ({file_name})")
            return None
            
    except Exception as e:
        st.error(f"Error reading file {uploaded_file.name}: {e}")
        return None

def build_gold_standard_prompt(en_files, pt_files):
    """Creates a few-shot prompt string from the gold standard files."""
    prompt = ""
    if not en_files or not pt_files:
        return prompt

    if len(en_files) != len(pt_files):
        st.sidebar.warning("Warning: The number of English and Portuguese gold standard files does not match. Using the minimum common number.")

    prompt = "\n\nHere are some 'gold standard' examples of English-to-Portuguese translations to guide your tone and terminology. Follow these examples closely:\n"
    
    for i, (en_file, pt_file) in enumerate(zip(en_files, pt_files)):
        en_text = read_file(en_file)
        pt_text = read_file(pt_file)
        
        if en_text and pt_text:
            prompt += f"\n--- Gold Standard Example {i+1} ---\n"
            prompt += f"[English]:\n{en_text}\n"
            prompt += f"[Portuguese]:\n{pt_text}\n"
            prompt += "--- End Example ---\n"
            
    return prompt

# --- MODIFIED FUNCTION ---
def call_gemini(api_key, prompt, task_description):
    """Generic function to call the Gemini API with error handling."""
    try:
        genai.configure(api_key=api_key)
        
        # --- FIX: Updated model name from 'gemini-1.5-flash' to 'gemini-2.5-flash' ---
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        with st.spinner(f"Gemini is {task_description}..."):
            response = model.generate_content(prompt, request_options={'timeout': 600})
            return response.text
    except Exception as e:
        if "API_KEY_INVALID" in str(e) or "PERMISSION_DENIED" in str(e):
             st.error(f"Error: Invalid Gemini API Key. Please check your .env file.")
        else:
            # Display the full error to the user
            st.error(f"Error communicating with Gemini: {e}")
        return None

def create_word_document(text_content):
    """Creates a .docx file in memory from a string."""
    doc = docx.Document()
    for para in text_content.split('\n'):
        doc.add_paragraph(para)
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

# --- Session State Initialization ---
if "project_started" not in st.session_state:
    st.session_state.project_started = False
    st.session_state.api_key = os.getenv("GEMINI_API_KEY") 
    st.session_state.source_text = None
    st.session_state.gold_standard_prompt = ""
    st.session_state.translation_step_4 = None
    st.session_state.translation_step_5 = None
    st.session_state.translation_step_6 = None
    st.session_state.final_text = None

# --- Sidebar ---
with st.sidebar:
    st.header("⚙️ Configuration")
    
    if st.session_state.api_key:
        st.success("✅ Gemini API Key loaded from .env")
    else:
        st.error("❌ Gemini API Key not found.")
        st.info("Please create a `.env` file in the app directory and add: `GEMINI_API_KEY='your_key_here'`")

    st.markdown("---")
    st.header("🥇 Gold Standard Samples")
    st.caption("Upload paired EN/PT files to be used as examples for the AI, improving terminology and style consistency.")
    
    gold_en_files = st.file_uploader(
        "Upload English (EN) Files", 
        type=["txt", "pdf", "docx"], 
        accept_multiple_files=True,
        key="gold_en"
    )
    gold_pt_files = st.file_uploader(
        "Upload Portuguese (PT) Files", 
        type=["txt", "pdf", "docx"], 
        accept_multiple_files=True,
        key="gold_pt"
    )

# --- Main App ---
st.title("🌐 Professional Translation Workflow Simulator")
st.markdown("This app simulates a 10-step translation process using Google Gemini for linguistic tasks.")
st.markdown("---")

# --- Step 1: Inquiry ---
st.header("1. Client Inquiry & Project Analysis")
col1, col2 = st.columns(2)
with col1:
    source_file = st.file_uploader(
        "Upload your source document", 
        type=["txt", "pdf", "docx"],
        help="Upload the .txt, .pdf, or .docx file you want to translate."
    )
with col2:
    source_lang = st.selectbox("Source Language", ["English", "Portuguese", "Spanish", "French", "German"])
    target_lang = st.selectbox("Target Language", ["Portuguese", "English", "Spanish", "French", "German"])

start_button = st.button("🚀 Start Project & Analyze", type="primary")

# --- Workflow Execution ---
if start_button or st.session_state.project_started:
    
    if not st.session_state.api_key:
        st.error("❌ Please provide your Gemini API Key in the `.env` file to begin.")
        st.stop()
        
    if not source_file:
        st.error("❌ Please upload a source file to translate.")
        st.stop()
        
    st.session_state.project_started = True

    # --- Step 2: Preparation ---
    with st.expander("2. Project Preparation", expanded=True):
        if st.session_state.source_text is None: 
            with st.spinner("Assigning PM, preparing files, setting up TM..."):
                st.session_state.source_text = read_file(source_file)
                st.session_state.gold_standard_prompt = build_gold_standard_prompt(gold_en_files, gold_pt_files)
                
                if st.session_state.source_text:
                    st.success("✔️ Files prepared. Project Manager assigned.")
                    word_count = len(st.session_state.source_text.split())
                    st.metric("Project Scope", f"{word_count} words")
                    if st.session_state.gold_standard_prompt:
                        st.info("✔️ Gold standard samples have been loaded and will be used.")
                else:
                    st.error("Failed to read source file. Please check the file format.")
                    st.stop()
        
        st.text_area("Extracted Source Text (Preview)", st.session_state.source_text, height=150, disabled=True)

    # --- Step 3: Translator Selection ---
    with st.expander("3. Translator Selection"):
        st.info("🤖 Qualified native linguist (Gemini 2.5 Flash) has been assigned based on subject matter.")

    # --- Step 4: Translation ---
    with st.expander("4. Translation Phase", expanded=True):
        if st.button("Run Translation (Step 4)") or st.session_state.translation_step_4:
            if not st.session_state.translation_step_4: 
                prompt = f"""You are a professional {source_lang}-to-{target_lang} translator.
                Translate the following text. Maintain a professional tone and ensure accuracy.
                Preserve paragraph breaks (indicated by newlines).
                {st.session_state.gold_standard_prompt}
                ---
                Source Text to Translate:
                {st.session_state.source_text}
                ---
                {target_lang} Translation:"""
                
                translation = call_gemini(st.session_state.api_key, prompt, "translating")
                if translation:
                    st.session_state.translation_step_4 = translation
                    st.session_state.final_text = translation 
                    st.success("Translation complete.")
                
            if st.session_state.translation_step_4:
                st.text_area("Initial Translation (from Gemini)", st.session_state.translation_step_4, height=200)

    # --- Step 5: Editing ---
    with st.expander("5. Editing (Second Linguist Review)"):
        if st.session_state.translation_step_4:
            if st.button("🤖 Ask Gemini to Edit/Review (Step 5)"):
                prompt = f"""You are a professional editor. Review the following translation from {source_lang} to {target_lang}.
                Compare it against the source text for accuracy, terminology, and tone.
                Correct any stylistic or grammatical issues to improve fluency. Preserve paragraph breaks.
                {st.session_state.gold_standard_prompt}
                ---
                Source Text:
                {st.session_state.source_text}
                ---
                Initial Translation to Review:
                {st.session_state.translation_step_4}
                ---
                Provide only the final, improved {target_lang} translation:"""
                
                edited_translation = call_gemini(st.session_state.api_key, prompt, "editing")
                if edited_translation:
                    st.session_state.translation_step_5 = edited_translation
                    st.session_state.final_text = edited_translation 
                    st.success("Edit complete.")
            
            default_text = st.session_state.translation_step_5 or st.session_state.translation_step_4
            manual_edit = st.text_area("Manually Edit Translation:", default_text, height=200, key="manual_edit")
            
            if manual_edit != default_text:
                st.session_state.translation_step_5 = manual_edit
                st.session_state.final_text = manual_edit
                st.info("Manual edit saved.")
        else:
            st.warning("Please complete Step 4 (Translation) first.")

    # --- Step 6: Proofreading / QA ---
    with st.expander("6. Proofreading / QA"):
        current_text_for_proofread = st.session_state.translation_step_5 or st.session_state.translation_step_4
        
        if current_text_for_proofread:
            if st.button("🤖 Ask Gemini for Final Proofread (Step 6)"):
                prompt = f"""You are a meticulous proofreader. Perform a final check on the following {target_lang} text.
                Correct only objective errors (typos, grammar, punctuation). Preserve paragraph breaks.
                Do NOT change the style or word choice unless it's grammatically incorrect.
                ---
                Text to Proofread:
                {current_text_for_proofread}
                ---
                Provide only the final, proofread text:"""
                
                proofread_text = call_gemini(st.session_state.api_key, prompt, "proofreading")
                if proofread_text:
                    st.session_state.translation_step_6 = proofread_text
                    st.session_state.final_text = proofread_text 
                    st.success("Proofreading complete.")

            if st.session_state.translation_step_6:
                 st.text_area("Final Proofread Text", st.session_state.translation_step_6, height=200, disabled=True)
        else:
            st.warning("Please complete Step 4 or 5 first.")

    # --- Steps 7-9: DTP, QC, Delivery ---
    if st.session_state.final_text:
        with st.expander("7. Desktop Publishing (DTP)"):
            st.info("ℹ️ DTP would occur here. This app delivers a clean .docx file, but manual layout adjustments would be needed for complex formats.")

        with st.expander("8. Final Quality Control"):
            st.success("✔️ PM final check complete. Files, naming conventions, and deliverables match client specifications.")

        with st.expander("9. Delivery", expanded=True):
            st.subheader("🎉 Final Deliverable Ready")
            st.markdown(f"The final translated text ({target_lang}) is below:")
            
            st.text_area("Final Text", st.session_state.final_text, height=300, disabled=True)
            
            try:
                # Create the Word document in memory
                doc_data = create_word_document(st.session_state.final_text)
                
                # Create a file name for the download
                base_name = os.path.splitext(source_file.name)[0]
                download_file_name = f"translated_{base_name}.docx"
                
                st.download_button(
                    label="⬇️ Download Final Translation (.docx)",
                    data=doc_data,
                    file_name=download_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Error creating .docx file for download: {e}")

    # --- Step 10: Feedback ---
    with st.expander("10. Client Feedback & Archiving"):
        st.slider("Please rate this translation:", 1, 5, 4)
        st.text_area("Provide any feedback (optional):")
        
        if st.button("Submit Feedback & Archive Project"):
            st.success("Thank you for your feedback! The project has been securely archived. The TM and glossary have been updated.")
            
            api_key = st.session_state.api_key
            for key in st.session_state.keys():
                del st.session_state[key]
            st.session_state.api_key = api_key 
            st.rerun()