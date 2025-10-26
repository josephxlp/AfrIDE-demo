import gradio as gr
import google.generativeai as genai
import pdfplumber
import docx  # Used for reading and creating .docx files
import io
import os
import tempfile
from dotenv import load_dotenv

# --- Load environment variables ---
load_dotenv()

# --- Helper Function: Read File ---
def read_file(filepath):
    """Reads the content of a file (txt, pdf, docx) from a filepath."""
    if filepath is None:
        return None
    try:
        file_name = os.path.basename(filepath)

        if file_name.endswith(".txt"):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_name.endswith(".pdf"):
            with pdfplumber.open(filepath) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        
        elif file_name.endswith(".docx"):
            doc = docx.Document(filepath)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        else:
            raise gr.Error(f"Unsupported file format: {file_name}")
            
    except Exception as e:
        raise gr.Error(f"Error reading file {file_name}: {e}")

# --- Helper Function: Build Gold Standard Prompt ---
def build_gold_standard_prompt(en_files, pt_files):
    """Creates a few-shot prompt string from the gold standard file objects."""
    prompt = ""
    if not en_files or not pt_files:
        return prompt

    if len(en_files) != len(pt_files):
        gr.Warning("Warning: The number of English and Portuguese gold standard files does not match.")

    prompt = "\n\nHere are some 'gold standard' examples of English-to-Portuguese translations to guide your tone and terminology. Follow these examples closely:\n"
    
    for i, (en_file, pt_file) in enumerate(zip(en_files, pt_files)):
        en_text = read_file(en_file.name) # .name is the path to the temp file
        pt_text = read_file(pt_file.name)
        
        if en_text and pt_text:
            prompt += f"\n--- Gold Standard Example {i+1} ---\n"
            prompt += f"[English]:\n{en_text}\n"
            prompt += f"[Portuguese]:\n{pt_text}\n"
            prompt += "--- End Example ---\n"
            
    return prompt

# --- Helper Function: Call Gemini API ---
def call_gemini(api_key, model_name, prompt, task_description):
    """Generic function to call the Gemini API with error handling."""
    if not api_key:
        raise gr.Error("Error: Invalid Gemini API Key. Please check your .env file.")
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt, request_options={'timeout': 600})
        return response.text
    except Exception as e:
        if "API_KEY_INVALID" in str(e) or "PERMISSION_DENIED" in str(e):
             raise gr.Error("Error: Invalid Gemini API Key. Please check your .env file.")
        else:
             raise gr.Error(f"Error communicating with Gemini: {e}")

# --- Helper Function: Create Word Doc ---
def create_word_document(text_content):
    """Creates a .docx file in memory from a string."""
    doc = docx.Document()
    if text_content:
        for para in text_content.split('\n'):
            doc.add_paragraph(para)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

# --- Prompt Generation Helpers ---
def generate_step_4_prompt(source_lang, target_lang, gold_prompt, source_text):
    return f"""You are a professional {source_lang}-to-{target_lang} translator.
Translate the following text. Maintain a professional tone and ensure accuracy.
Preserve paragraph breaks (indicated by newlines).
{gold_prompt}
---
Source Text to Translate:
{source_text}
---
{target_lang} Translation:"""

def generate_step_5_prompt(source_lang, target_lang, gold_prompt, source_text, translation_text):
    return f"""You are a professional editor. Review the following translation from {source_lang} to {target_lang}.
Compare it against the source text for accuracy, terminology, and tone.
Correct any stylistic or grammatical issues to improve fluency. Preserve paragraph breaks.
{gold_prompt}
---
Source Text:
{source_text}
---
Initial Translation to Review:
{translation_text}
---
Provide only the final, improved {target_lang} translation:"""

def generate_step_6_prompt(target_lang, text_to_proofread):
    return f"""You are a meticulous proofreader. Perform a final check on the following {target_lang} text.
Correct only objective errors (typos, grammar, punctuation). Preserve paragraph breaks.
Do NOT change the style or word choice unless it's grammatically incorrect.
---
Text to Proofread:
{text_to_proofread}
---
Provide only the final, proofread text:"""


# --- Gradio Event Handlers ---

def start_project(source_file, en_files, pt_files, source_lang, target_lang):
    """Handles the 'Start Project' button click."""
    
    # 1. Check API Key and Source File
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise gr.Error("❌ Gemini API Key not found in .env file.")
    if not source_file:
        raise gr.Error("❌ Please upload a source file to translate.")

    # 2. Read Files
    source_text = read_file(source_file.name)
    gold_prompt = build_gold_standard_prompt(en_files, pt_files)

    if not source_text:
        raise gr.Error("Failed to read source file.")

    # 3. Generate Stats and Prompts
    word_count = len(source_text.split())
    word_count_md = f"**Project Scope:** {word_count} words"
    gold_status_md = "✔️ Gold standard samples have been loaded." if gold_prompt else "ℹ️ No gold standard samples loaded."
    
    prompt_4 = generate_step_4_prompt(source_lang, target_lang, gold_prompt, source_text)
    
    # 4. Return dictionary to update all UI components
    return {
        # Update State
        api_key_state: api_key,
        source_text_state: source_text,
        gold_prompt_state: gold_prompt,
        source_file_obj_state: source_file,
        
        # Update Step 2 (Preparation)
        word_count_label: word_count_md,
        gold_status_label: gold_status_md,
        source_text_widget: source_text,
        
        # Update Step 4
        step_4_prompt_text: prompt_4,
        step_4_source_text: source_text,
        
        # Update Step 5
        step_5_source_text: source_text,
        
        # Update Step 6
        step_6_source_text: source_text,

        # Update Step 9
        final_source_widget: source_text,
        
        # Make workflow visible
        step_2_accordion: gr.Accordion(visible=True),
        step_3_accordion: gr.Accordion(visible=True),
        step_4_accordion: gr.Accordion(visible=True),
        step_5_accordion: gr.Accordion(visible=True),
        step_6_accordion: gr.Accordion(visible=True),
        step_7_accordion: gr.Accordion(visible=True),
        step_8_accordion: gr.Accordion(visible=True),
        step_9_accordion: gr.Accordion(visible=True),
        step_10_accordion: gr.Accordion(visible=True),
        
        # Disable Step 1
        source_file_upload: gr.File(interactive=False),
        gold_en_upload: gr.File(interactive=False),
        gold_pt_upload: gr.File(interactive=False),
        source_lang_dd: gr.Dropdown(interactive=False),
        target_lang_dd: gr.Dropdown(interactive=False),
        start_button: gr.Button(interactive=False),
    }

def run_step_4(prompt_4, model_name, api_key, source_lang, target_lang, source_text, gold_prompt):
    """Handles the 'Run Translation (Step 4)' button click."""
    translation = call_gemini(api_key, model_name, prompt_4, "translating")
    
    if translation:
        # Generate next prompt
        prompt_5 = generate_step_5_prompt(source_lang, target_lang, gold_prompt, source_text, translation)
        
        return {
            translation_step_4_state: translation,
            final_text_state: translation,
            step_4_target_text: translation,
            step_5_target_text: translation,
            final_translation_widget: translation,
            step_5_prompt_text: prompt_5, # Update step 5 prompt
        }
    return {} # No update on failure

def run_step_5_ai(prompt_5, model_name, api_key, target_lang, final_text_state):
    """Handles the 'Ask Gemini to Edit/Review (Step 5)' button click."""
    edited_translation = call_gemini(api_key, model_name, prompt_5, "editing")
    
    if edited_translation:
        # Generate next prompt
        prompt_6 = generate_step_6_prompt(target_lang, edited_translation)
        
        return {
            translation_step_5_state: edited_translation,
            final_text_state: edited_translation,
            step_5_target_text: edited_translation,
            final_translation_widget: edited_translation,
            step_6_prompt_text: prompt_6, # Update step 6 prompt
        }
    return {}

def on_manual_edit(manual_text, target_lang, final_text_state):
    """Handles live manual editing in Step 5."""
    # Generate next prompt based on the manual edit
    prompt_6 = generate_step_6_prompt(target_lang, manual_text)
    
    return {
        final_text_state: manual_text,
        final_translation_widget: manual_text,
        step_6_prompt_text: prompt_6, # Update step 6 prompt in real-time
    }

def run_step_6(prompt_6, model_name, api_key):
    """Handles the 'Final Proofread (Step 6)' button click."""
    proofread_text = call_gemini(api_key, model_name, prompt_6, "proofreading")
    
    if proofread_text:
        return {
            translation_step_6_state: proofread_text,
            final_text_state: proofread_text,
            step_6_target_text: proofread_text,
            final_translation_widget: proofread_text,
        }
    return {}

def download_docx(final_text, source_file_obj):
    """Creates the .docx file and returns its path for download."""
    if not final_text:
        gr.Warning("No final text to download.")
        return None
        
    try:
        base_name = os.path.splitext(os.path.basename(source_file_obj.name))[0]
        download_file_name = f"translated_{base_name}.docx"
        
        doc_data = create_word_document(final_text)
        
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix=download_file_name) as temp_f:
            temp_f.write(doc_data)
            temp_path = temp_f.name
            
        gr.Info("Download file prepared.")
        return { download_file_widget: gr.File(value=temp_path, label="Download Final Translation (.docx)") }

    except Exception as e:
        raise gr.Error(f"Error creating .docx file for download: {e}")

def archive_project():
    """Resets the entire UI to its initial state."""
    gr.Info("Project archived. Ready for new project.")
    return {
        # Reset State
        api_key_state: None,
        source_text_state: None,
        gold_prompt_state: "",
        source_file_obj_state: None,
        translation_step_4_state: None,
        translation_step_5_state: None,
        translation_step_6_state: None,
        final_text_state: None,
        
        # Reset Step 1 (and re-enable)
        source_file_upload: gr.File(value=None, interactive=True),
        gold_en_upload: gr.File(value=None, interactive=True),
        gold_pt_upload: gr.File(value=None, interactive=True),
        source_lang_dd: gr.Dropdown(value="English", interactive=True),
        target_lang_dd: gr.Dropdown(value="Portuguese", interactive=True),
        start_button: gr.Button(interactive=True),

        # Reset Step 2
        word_count_label: "",
        gold_status_label: "",
        source_text_widget: "",
        
        # Reset Step 4
        step_4_prompt_text: "",
        step_4_source_text: "",
        step_4_target_text: "",
        
        # Reset Step 5
        step_5_prompt_text: "",
        step_5_source_text: "",
        step_5_target_text: "",
        
        # Reset Step 6
        step_6_prompt_text: "",
        step_6_source_text: "",
        step_6_target_text: "",

        # Reset Step 9
        final_source_widget: "",
        final_translation_widget: "",
        download_file_widget: gr.File(value=None, label="Download Final Translation (.docx)"),
        
        # Reset Step 10
        feedback_slider: gr.Slider(value=4),
        feedback_text: "",

        # Hide workflow
        step_2_accordion: gr.Accordion(visible=False),
        step_3_accordion: gr.Accordion(visible=False),
        step_4_accordion: gr.Accordion(visible=False),
        step_5_accordion: gr.Accordion(visible=False),
        step_6_accordion: gr.Accordion(visible=False),
        step_7_accordion: gr.Accordion(visible=False),
        step_8_accordion: gr.Accordion(visible=False),
        step_9_accordion: gr.Accordion(visible=False),
        step_10_accordion: gr.Accordion(visible=False),
    }

# --- Gradio UI Layout ---

with gr.Blocks(theme=gr.themes.Soft(), title="Professional Translation Workflow") as demo:
    
    # --- Define State Variables ---
    api_key_state = gr.State(None)
    source_text_state = gr.State(None)
    gold_prompt_state = gr.State("")
    source_file_obj_state = gr.State(None)
    translation_step_4_state = gr.State(None)
    translation_step_5_state = gr.State(None)
    translation_step_6_state = gr.State(None)
    final_text_state = gr.State(None)

    with gr.Row():
        # --- Sidebar ---
        with gr.Column(scale=1, min_width=350):
            with gr.Group():
                gr.Markdown("## ⚙️ Configuration")
                api_key_status = gr.Markdown("✅ Gemini API Key loaded from .env" if os.getenv("GEMINI_API_KEY") else "❌ Gemini API Key not found in .env")
                
                gr.Markdown("### Select Model")
                model_name_dd = gr.Dropdown(
                    ['gemini-1.5-flash-latest', 'gemini-1.5-pro-latest'], 
                    label="Gemini Model", 
                    value='gemini-1.5-flash-latest'
                )

            with gr.Group():
                gr.Markdown("## 🥇 Gold Standard Samples")
                gr.Markdown("Upload paired EN/PT files for examples.")
                gold_en_upload = gr.File(label="Upload English (EN) Files", file_count="multiple", file_types=[".txt", ".pdf", ".docx"])
                gold_pt_upload = gr.File(label="Upload Portuguese (PT) Files", file_count="multiple", file_types=[".txt", ".pdf", ".docx"])

        # --- Main Content ---
        with gr.Column(scale=3):
            gr.Markdown("# 🌐 Professional Translation Workflow Simulator")
            gr.Markdown("This app simulates a 10-step translation process using Google Gemini.")
            
            # --- Step 1: Inquiry ---
            with gr.Accordion("1. Client Inquiry & Project Analysis", open=True):
                with gr.Row():
                    with gr.Column():
                        source_file_upload = gr.File(label="Upload your source document", file_types=[".txt", ".pdf", ".docx"])
                    with gr.Column():
                        lang_list = ["English", "Portuguese", "Spanish", "French", "German"]
                        source_lang_dd = gr.Dropdown(lang_list, label="Source Language", value="English")
                        target_lang_dd = gr.Dropdown(lang_list, label="Target Language", value="Portuguese")
                start_button = gr.Button("🚀 Start Project & Analyze", variant="primary")
            
            # --- Steps 2-10 (Initially Hidden) ---
            with gr.Accordion("2. Project Preparation", open=True, visible=False) as step_2_accordion:
                word_count_label = gr.Markdown("**Project Scope:** 0 words")
                gold_status_label = gr.Markdown("ℹ️ No gold standard samples loaded.")
                gr.Markdown("### Extracted Source Text (Preview)")
                source_text_widget = gr.Textbox(label="Source Text", lines=8, interactive=False)

            with gr.Accordion("3. Translator Selection", visible=False) as step_3_accordion:
                gr.Info("🤖 Qualified native linguist (Gemini) has been assigned.")

            with gr.Accordion("4. Translation Phase", visible=False) as step_4_accordion:
                step_4_prompt_text = gr.Textbox(label="Step 4 Prompt (Editable)", lines=8, interactive=True)
                step_4_button = gr.Button("Run Translation (Step 4)", variant="secondary")
                with gr.Row():
                    step_4_source_text = gr.Textbox(label="Source Text", lines=10, interactive=False)
                    step_4_target_text = gr.Textbox(label="Initial Translation (from Gemini)", lines=10, interactive=False)

            with gr.Accordion("5. Editing (Second Linguist Review)", visible=False) as step_5_accordion:
                step_5_prompt_text = gr.Textbox(label="Step 5 Prompt (Editable)", lines=8, interactive=True)
                step_5_button = gr.Button("🤖 Ask Gemini to Edit/Review (Step 5)", variant="secondary")
                with gr.Row():
                    step_5_source_text = gr.Textbox(label="Source Text", lines=10, interactive=False)
                    step_5_target_text = gr.Textbox(label="Manually Edit Translation", lines=10, interactive=True) # <-- MANUAL EDITING

            with gr.Accordion("6. Proofreading / QA", visible=False) as step_6_accordion:
                step_6_prompt_text = gr.Textbox(label="Step 6 Prompt (Editable)", lines=8, interactive=True)
                step_6_button = gr.Button("🤖 Ask Gemini for Final Proofread (Step 6)", variant="secondary")
                with gr.Row():
                    step_6_source_text = gr.Textbox(label="Source Text", lines=10, interactive=False)
                    step_6_target_text = gr.Textbox(label="Final Proofread Text", lines=10, interactive=False)

            with gr.Accordion("7. Desktop Publishing (DTP)", visible=False) as step_7_accordion:
                gr.Info("ℹ️ DTP would occur here. This app delivers a clean .docx file.")

            with gr.Accordion("8. Final Quality Control", visible=False) as step_8_accordion:
                gr.Success("✔️ PM final check complete.")

            with gr.Accordion("9. Delivery", visible=False) as step_9_accordion:
                gr.Markdown("## 🎉 Final Deliverable Ready")
                with gr.Row():
                    final_source_widget = gr.Textbox(label="Final Source Text", lines=12, interactive=False)
                    final_translation_widget = gr.Textbox(label="Final Translation", lines=12, interactive=False)
                
                prepare_download_button = gr.Button("⬇️ Prepare Download (.docx)")
                download_file_widget = gr.File(label="Download Final Translation (.docx)")

            with gr.Accordion("10. Client Feedback & Archiving", visible=False) as step_10_accordion:
                feedback_slider = gr.Slider(1, 5, value=4, step=1, label="Please rate this translation:")
                feedback_text = gr.Textbox(label="Provide any feedback (optional):", lines=3)
                archive_button = gr.Button("Submit Feedback & Archive Project")

    # --- Wire up Event Handlers ---
    
    # Step 1
    start_button.click(
        fn=start_project,
        inputs=[source_file_upload, gold_en_upload, gold_pt_upload, source_lang_dd, target_lang_dd],
        outputs=[
            api_key_state, source_text_state, gold_prompt_state, source_file_obj_state,
            word_count_label, gold_status_label, source_text_widget,
            step_4_prompt_text, step_4_source_text,
            step_5_source_text,
            step_6_source_text,
            final_source_widget,
            step_2_accordion, step_3_accordion, step_4_accordion, step_5_accordion,
            step_6_accordion, step_7_accordion, step_8_accordion, step_9_accordion, step_10_accordion,
            source_file_upload, gold_en_upload, gold_pt_upload, source_lang_dd, target_lang_dd, start_button
        ]
    )
    
    # Step 4
    step_4_button.click(
        fn=run_step_4,
        inputs=[
            step_4_prompt_text, model_name_dd, api_key_state, 
            source_lang_dd, target_lang_dd, source_text_state, gold_prompt_state
        ],
        outputs=[
            translation_step_4_state, final_text_state,
            step_4_target_text, step_5_target_text, final_translation_widget,
            step_5_prompt_text
        ]
    )

    # Step 5 (AI)
    step_5_button.click(
        fn=run_step_5_ai,
        inputs=[step_5_prompt_text, model_name_dd, api_key_state, target_lang_dd, final_text_state],
        outputs=[
            translation_step_5_state, final_text_state,
            step_5_target_text, final_translation_widget,
            step_6_prompt_text
        ]
    )
    
    # Step 5 (Manual Edit)
    step_5_target_text.input(
        fn=on_manual_edit,
        inputs=[step_5_target_text, target_lang_dd, final_text_state],
        outputs=[final_text_state, final_translation_widget, step_6_prompt_text]
    )
    
    # Step 6
    step_6_button.click(
        fn=run_step_6,
        inputs=[step_6_prompt_text, model_name_dd, api_key_state],
        outputs=[
            translation_step_6_state, final_text_state,
            step_6_target_text, final_translation_widget
        ]
    )
    
    # Step 9
    prepare_download_button.click(
        fn=download_docx,
        inputs=[final_text_state, source_file_obj_state],
        outputs=[download_file_widget]
    )

    # Step 10
    archive_button.click(
        fn=archive_project,
        inputs=None,
        outputs=[
            # State
            api_key_state, source_text_state, gold_prompt_state, source_file_obj_state,
            translation_step_4_state, translation_step_5_state, translation_step_6_state, final_text_state,
            # Step 1
            source_file_upload, gold_en_upload, gold_pt_upload, source_lang_dd, target_lang_dd, start_button,
            # Step 2
            word_count_label, gold_status_label, source_text_widget,
            # Step 4
            step_4_prompt_text, step_4_source_text, step_4_target_text,
            # Step 5
            step_5_prompt_text, step_5_source_text, step_5_target_text,
            # Step 6
            step_6_prompt_text, step_6_source_text, step_6_target_text,
            # Step 9
            final_source_widget, final_translation_widget, download_file_widget,
            # Step 10
            feedback_slider, feedback_text,
            # Accordions
            step_2_accordion, step_3_accordion, step_4_accordion, step_5_accordion,
            step_6_accordion, step_7_accordion, step_8_accordion, step_9_accordion, step_10_accordion,
        ]
    )

if __name__ == "__main__":
    demo.launch()